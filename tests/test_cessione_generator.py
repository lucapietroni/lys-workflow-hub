"""Test del workflow Cessione del Credito: CF parser, builder, generatore docx."""
from __future__ import annotations

from datetime import date, datetime
from io import BytesIO
from unittest.mock import MagicMock

from docx import Document

from lys_workflow_hub.core.wincar_repository import (
    Cliente,
    CompagniaCliente,
    Controparte,
    Pratica,
    Sinistro,
    Veicolo,
)
from lys_workflow_hub.workflows.cessione_credito import (
    CessioneData,
    filename_for,
    from_pratica,
    generate,
)
from lys_workflow_hub.workflows.cessione_credito.cf_parser import (
    parse_codice_fiscale,
)


# ---------------------------------------------------------------------------
# CF parser
# ---------------------------------------------------------------------------


def test_cf_parser_decodes_valid_male_cf():
    info = parse_codice_fiscale("RSSMRA80A01H501U")
    assert info.valido is True
    assert info.data_nascita == date(1980, 1, 1)
    assert info.luogo_nascita is not None
    assert "Roma" in info.luogo_nascita
    assert info.sesso == "M"
    assert info.articolo == "Il"
    assert info.desinenza_nato == "o"


def test_cf_parser_decodes_valid_female_cf():
    """Genera dinamicamente un CF femminile valido per non dipendere dai check digit."""
    from codicefiscale import codicefiscale as cf_lib
    cf = cf_lib.encode(
        lastname="Rossi", firstname="Maria", gender="F",
        birthdate="1985-04-01", birthplace="Roma",
    )
    info = parse_codice_fiscale(cf)
    assert info.valido is True
    assert info.sesso == "F"
    assert info.articolo == "La"
    assert info.desinenza_nato == "a"


def test_cf_parser_rejects_partita_iva():
    info = parse_codice_fiscale("13633531002")
    assert info.valido is False
    assert info.motivo_invalidita is not None and "partita" in info.motivo_invalidita.lower()


def test_cf_parser_rejects_empty():
    assert parse_codice_fiscale(None).valido is False
    assert parse_codice_fiscale("").valido is False
    assert parse_codice_fiscale("   ").valido is False


def test_cf_parser_rejects_bad_length():
    info = parse_codice_fiscale("ABC123")
    assert info.valido is False
    assert "lunghezza" in info.motivo_invalidita.lower()


# ---------------------------------------------------------------------------
# Builder from_pratica
# ---------------------------------------------------------------------------


def _full_pratica(cf: str = "RSSMRA80A01H501U", piva: str = "") -> Pratica:
    return Pratica(
        numero=766,
        data_creazione=datetime(2026, 5, 10, 14, 30),
        cliente=Cliente(
            nominativo="ROSSI MARIO",
            codice_fiscale=cf,
            partita_iva=piva,
            via="Via Roma 12",
            citta="Roma",
            cap="00100",
            provincia="RM",
            telefono=None, cellulare=None, email=None,
        ),
        veicolo=Veicolo(targa="AB123CD", marca="FIAT", modello="Punto 1.2", telaio=None),
        sinistro=Sinistro(
            data=date(2026, 5, 8), ora="10:15",
            comune="Roma", via="Via Nazionale",
            dinamica="Tamponamento al semaforo.",
            numero=None, tipo=None,
        ),
        controparte=Controparte(
            proprietario="BIANCHI LUCA", conducente="BIANCHI LUCA",
            veicolo_descrizione="BMW Serie 1", targa="XY987ZW",
            indirizzo=None, citta=None,
            compagnia="Generali Italia SpA", numero_polizza="POL-99887766",
        ),
        assicurazione_cliente=CompagniaCliente(
            nome="Allianz", indirizzo=None, citta=None, cap=None, provincia=None,
            numero_polizza="POL-11223344", agenzia=None,
        ),
    )


def test_from_pratica_basic():
    p = _full_pratica()
    data = from_pratica(p)
    assert isinstance(data, CessioneData)
    assert data.numero_pratica == 766
    assert data.cedente_nome_completo == "ROSSI MARIO"
    assert data.cedente_codice_fiscale == "RSSMRA80A01H501U"
    assert data.cedente_data_nascita == date(1980, 1, 1)
    assert "Roma" in data.cedente_luogo_nascita
    assert data.cedente_sesso == "M"
    assert data.e_ditta is False


def test_from_pratica_marks_ditta_when_partita_iva():
    p = _full_pratica(cf="", piva="13633531002")
    data = from_pratica(p)
    assert data.e_ditta is True
    assert data.ditta_nome == "ROSSI MARIO"
    assert data.ditta_partita_iva == "13633531002"


def test_from_pratica_applies_overrides():
    p = _full_pratica()
    data = from_pratica(p, overrides={
        "sinistro_dinamica": "DINAMICA SOSTITUITA DALL'OPERATORE",
        "cedente_residenza_via": "Via Override 1",
    })
    assert data.sinistro_dinamica == "DINAMICA SOSTITUITA DALL'OPERATORE"
    assert data.cedente_residenza_via == "Via Override 1"


def test_campi_mancanti_returns_empty_when_complete():
    p = _full_pratica()
    data = from_pratica(p)
    assert data.campi_mancanti() == []


def test_campi_mancanti_lists_missing():
    p = _full_pratica()
    data = from_pratica(p, overrides={"sinistro_dinamica": "", "controparte_polizza": ""})
    missing = data.campi_mancanti()
    assert "Dinamica del sinistro" in missing
    assert "Numero polizza controparte" in missing


# ---------------------------------------------------------------------------
# Generator
# ---------------------------------------------------------------------------


def test_generate_produces_valid_docx_with_expected_content():
    data = from_pratica(_full_pratica())
    docx_bytes = generate(data)
    assert docx_bytes[:4] == b"PK\x03\x04"  # zip magic = valido docx
    # Riapriamo e cerchiamo i campi chiave nel testo
    doc = Document(BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "CESSIONE DI CREDITO" in text
    assert "PREMESSO" in text
    assert "DICHIARA" in text
    assert "ROSSI MARIO" in text
    assert "RSSMRA80A01H501U" in text
    assert "Tamponamento al semaforo." in text
    assert "Generali Italia SpA" in text
    assert "AB123CD" in text
    assert "BMW Serie 1" in text
    assert "Carrozzeria LYS Auto srl" in text


def test_generate_respects_sex_articles():
    from codicefiscale import codicefiscale as cf_lib
    cf_femmina = cf_lib.encode(
        lastname="Rossi", firstname="Maria", gender="F",
        birthdate="1985-04-01", birthplace="Roma",
    )
    p = _full_pratica(cf=cf_femmina)
    data = from_pratica(p)
    docx_bytes = generate(data)
    text = "\n".join(p.text for p in Document(BytesIO(docx_bytes)).paragraphs)
    # Per donna l'incipit e' "La sottoscritta" e "nata a"
    assert "La sottoscritta" in text
    assert "nata a" in text


def test_generate_male_uses_masculine_form():
    p = _full_pratica()  # default CF maschile
    data = from_pratica(p)
    text = "\n".join(p.text for p in Document(BytesIO(generate(data))).paragraphs)
    assert "Il sottoscritto" in text
    assert "nato a" in text


def test_filename_includes_numero_and_nome():
    p = _full_pratica()
    data = from_pratica(p)
    fname = filename_for(data)
    assert fname.startswith("Cessione_credito_766_")
    assert fname.endswith(".docx")
    assert "ROSSI" in fname.upper()
