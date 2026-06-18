"""Generatore DOCX per verbali di consegna/riconsegna veicolo di cortesia."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from lys_workflow_hub.workflows.verbale_cortesia.data import (
    TIPO_USCITA,
    VerbaleData,
)
from lys_workflow_hub.workflows.cessione_credito.data import (
    CARROZZERIA_NOME,
    CARROZZERIA_VIA,
    CARROZZERIA_CAP,
    CARROZZERIA_COMUNE,
    CARROZZERIA_PIVA,
)


_ASSETS_DIR = Path(__file__).parent / "assets"
_LOGO_PNG = _ASSETS_DIR / "logo_lys.png"
_TIMBRO_PNG = _ASSETS_DIR / "timbro_lys.png"

COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_GREY = RGBColor(0x4A, 0x55, 0x68)
COLOR_HEADER_BG = "2C3E50"   # sfondo intestazione sezione (quasi-nero)
COLOR_SUBHDR_BG = "D0D0D0"   # sfondo sotto-intestazione colonne

# 17.6 cm in twips (A4 - margini 1.7cm × 2) — forza larghezza tabelle
TABLE_WIDTH_DXA = 9977


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------


def _set_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin   = Cm(1.7)
        section.right_margin  = Cm(1.7)


def _run(para, text: str, *, bold: bool = False, size: float | None = None,
         color: RGBColor | None = None, italic: bool = False) -> None:
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    if size  is not None: r.font.size      = Pt(size)
    if color is not None: r.font.color.rgb = color


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def _set_cell_margins(cell, top: int = 40, bottom: int = 40,
                      left: int = 80, right: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tc_pr.append(tcMar)


def _para(doc: Document, *, alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before: float = 0, space_after: float = 0) -> object:
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p


def _set_table_width(table) -> None:
    """Forza la tabella a TABLE_WIDTH_DXA twips, disabilita autofit."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for child in list(tbl_pr):
        if child.tag == qn("w:tblW"):
            tbl_pr.remove(child)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(TABLE_WIDTH_DXA))
    tblW.set(qn("w:type"), "dxa")
    tbl_pr.append(tblW)


def _set_row_height(table, row_idx: int, twips: int,
                    rule: str = "atLeast") -> None:
    tr_pr = table.rows[row_idx]._tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(twips))
    trH.set(qn("w:hRule"), rule)
    tr_pr.append(trH)


def _force_table_borders(table) -> None:
    """Imposta bordi espliciti — fix bordo superiore mancante con alcuni renderer."""
    tblPr = table._tbl.get_or_add_tblPr()
    for child in list(tblPr):
        if child.tag == qn("w:tblBorders"):
            tblPr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)


def _merge_row(table, row_idx: int, start_col: int, end_col: int):
    """Merge celle [start_col, end_col] nella riga row_idx, ritorna la cella."""
    row  = table.rows[row_idx]
    cell = row.cells[start_col]
    for i in range(start_col + 1, end_col + 1):
        cell = cell.merge(row.cells[i])
    return cell


def _lv(cell, label: str, value: str,
        label_size: float = 8.0, value_size: float = 9.0) -> None:
    _set_cell_margins(cell)
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    _run(para, f"{label}: ", bold=True, size=label_size, color=COLOR_GREY)
    _run(para, value or "",  size=value_size)


def _section_row(table, row_idx: int, text: str,
                 ncols: int, row_height: int = 330) -> None:
    """Prima riga della tabella come intestazione di sezione (sfondo scuro, testo bianco)."""
    cell = _merge_row(table, row_idx, 0, ncols - 1)
    _set_cell_shading(cell, COLOR_HEADER_BG)
    _set_cell_margins(cell, top=35, bottom=35, left=80, right=80)
    _set_row_height(table, row_idx, row_height, rule="exact")
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    _run(para, text.upper(), bold=True, size=9.0, color=COLOR_WHITE)


def _col_header_row(table, row_idx: int, labels: list[str],
                    row_height: int = 300) -> None:
    """Riga intestazione colonne (sfondo grigio chiaro)."""
    _set_row_height(table, row_idx, row_height, rule="exact")
    for i, (cell, label) in enumerate(zip(table.rows[row_idx].cells, labels)):
        _set_cell_shading(cell, COLOR_SUBHDR_BG)
        _set_cell_margins(cell, top=30, bottom=30, left=80, right=80)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, label, bold=True, size=8.5)


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------


def _add_header(doc: Document, data: VerbaleData) -> None:
    if _LOGO_PNG.exists():
        p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
        p.add_run().add_picture(str(_LOGO_PNG), width=Cm(5.0))

    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=1)
    _run(p,
         "VERBALE DI CONSEGNA VEICOLO DI CORTESIA"
         if data.tipo == TIPO_USCITA
         else "VERBALE DI RICONSEGNA VEICOLO DI CORTESIA",
         bold=True, size=13)

    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
    _run(p, f"(Verbale {data.label_tipo})  —  Pratica n° {data.numero_pratica}",
         size=9.0, italic=True, color=COLOR_GREY)


DATA_ROW_H = 330   # twips per riga dati standard


def _add_locatario(doc: Document, data: VerbaleData) -> None:
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    _set_table_width(table)

    _section_row(table, 0, "Locatario", 4)

    c = _merge_row(table, 1, 0, 1); _lv(c, "Locatario", data.locatario_nome)
    c = _merge_row(table, 1, 2, 3); _lv(c, "Cod. Fiscale", data.codice_fiscale)
    _set_row_height(table, 1, DATA_ROW_H)

    c = _merge_row(table, 2, 0, 1); _lv(c, "Indirizzo", data.indirizzo)
    _lv(table.rows[2].cells[2], "Località", data.localita)
    _lv(table.rows[2].cells[3], "CAP", data.cap)
    _set_row_height(table, 2, DATA_ROW_H)

    _lv(table.rows[3].cells[0], "Patente N°",    data.patente_numero)
    _lv(table.rows[3].cells[1], "Rilasciata da", data.patente_rilasciata_da)
    _lv(table.rows[3].cells[2], "il",            data.patente_data_rilascio)
    _lv(table.rows[3].cells[3], "Validità",      data.patente_validita)
    _set_row_height(table, 3, DATA_ROW_H)

    c = _merge_row(table, 4, 0, 3); _lv(c, "Telefono", data.telefono)
    _set_row_height(table, 4, DATA_ROW_H)


def _add_veicolo(doc: Document, data: VerbaleData) -> None:
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    _set_table_width(table)

    _section_row(table, 0, "Veicolo di cortesia", 4)

    c = _merge_row(table, 1, 0, 1); _lv(c, "Marca/Modello", data.marca_modello)
    _lv(table.rows[1].cells[2], "Telaio", data.telaio)
    _lv(table.rows[1].cells[3], "Targa",  data.targa)
    _set_row_height(table, 1, DATA_ROW_H)

    _lv(table.rows[2].cells[0], data.label_km,        data.km)
    _lv(table.rows[2].cells[1], "Livello carburante",  data.livello_carburante)
    c = _merge_row(table, 2, 2, 3); _lv(c, "Omologato per", data.omologato_per)
    _set_row_height(table, 2, DATA_ROW_H)

    _lv(table.rows[3].cells[0], "Max Km mese",         data.max_km_mese)
    _lv(table.rows[3].cells[1], "Max Km giorno",       data.max_km_giorno)
    c = _merge_row(table, 3, 2, 3); _lv(c, "Tariffa Km eccedenti", data.tariffa_km_eccedenti)
    _set_row_height(table, 3, DATA_ROW_H)

    c = _merge_row(table, 4, 0, 3); _lv(c, "Accessori", data.accessori)
    _set_row_height(table, 4, DATA_ROW_H + 80)


def _add_franchigie(doc: Document, data: VerbaleData) -> None:
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    _set_table_width(table)

    _section_row(table, 0, "Franchigie", 4)

    _lv(table.rows[1].cells[0], "RCA",           f"{data.rca} €" if data.rca else "")
    _lv(table.rows[1].cells[1], "KASCO",         f"{data.kasco} €" if data.kasco else "")
    c = _merge_row(table, 1, 2, 3); _lv(c, "Furto Incendio", data.furto_incendio)
    _set_row_height(table, 1, DATA_ROW_H)

    c = _merge_row(table, 2, 0, 3)
    _lv(c, "Importo giornaliero della locazione €", data.importo_giornaliero)
    _set_row_height(table, 2, DATA_ROW_H)


def _add_danni(doc: Document, data: VerbaleData) -> None:
    label = (
        "Distinta danni vettura"
        if data.tipo == TIPO_USCITA
        else "Distinta danni vettura alla riconsegna"
    )
    rows_data = list(data.danni) + [("", "")] * max(0, 3 - len(data.danni))
    n_rows = 2 + len(rows_data)   # section header + col header + dati

    table = doc.add_table(rows=n_rows, cols=2)
    table.style = "Table Grid"
    _set_table_width(table)

    _section_row(table, 0, label, 2)
    _col_header_row(table, 1, ["Parte danneggiata", "Dettaglio"])

    for i, (parte, det) in enumerate(rows_data, start=2):
        _set_row_height(table, i, DATA_ROW_H + 30)
        for cell, val in zip(table.rows[i].cells, [parte, det]):
            _set_cell_margins(cell)
            if val:
                _run(cell.paragraphs[0], val, size=9.0)

    p = _para(doc, space_before=2, space_after=2)
    _run(
        p,
        "Il Locatario dichiara di aver accertato le condizioni generali del veicolo "
        "e di aver riscontrato/non riscontrato le sue perfette condizioni.",
        italic=True, size=8.5,
    )


def _add_note(doc: Document, data: VerbaleData) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    _set_table_width(table)

    _section_row(table, 0, "Note", 1)

    cell = table.rows[1].cells[0]
    _set_cell_margins(cell, top=60, bottom=60)
    _set_row_height(table, 1, 550)
    if data.note:
        _run(cell.paragraphs[0], data.note, size=9.0)


def _add_firme(doc: Document, data: VerbaleData) -> None:
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    _set_table_width(table)

    _col_header_row(table, 0, [
        f"Data e ora di {data.label_tipo}",
        "Il Locatario  —  Firma",
        "Il Locatore  —  Timbro e firma",
    ], row_height=330)

    _set_row_height(table, 1, 1450)

    date_cell = table.rows[1].cells[0]
    _set_cell_margins(date_cell)
    date_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = date_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, data.data_ora, size=10.0)

    # Il Locatario — firma manuale cliente
    loc_cell = table.rows[1].cells[1]
    _set_cell_margins(loc_cell)
    loc_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p = loc_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "........................", size=9.0)

    # Il Locatore — timbro LYS Auto
    lys_cell = table.rows[1].cells[2]
    _set_cell_margins(lys_cell, top=20, bottom=20, left=20, right=20)
    lys_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = lys_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _TIMBRO_PNG.exists():
        p.add_run().add_picture(str(_TIMBRO_PNG), width=Cm(5.0))
    else:
        _run(p, "........................", size=9.0)


MOTIVAZIONI = [
    ("lavoro",      "MOTIVI DI LAVORO"),
    ("familiare",   "MOTIVI FAMILIARI"),
    ("unico_mezzo", "UNICO MEZZO A DISPOSIZIONE DEL NUCLEO FAMILIARE"),
    ("altro",       "ALTRO"),
]


def _add_dichiarazione(doc: Document, data: VerbaleData) -> None:
    """Pagina 2 — Dichiarazione di necessità auto sostitutiva (stesso stile pagina 1)."""
    doc.add_page_break()

    # --- Logo ---
    if _LOGO_PNG.exists():
        p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
        p.add_run().add_picture(str(_LOGO_PNG), width=Cm(4.5))

    # --- Titolo come riga tabella scura (1 col) ---
    t_title = doc.add_table(rows=1, cols=1)
    t_title.style = "Table Grid"
    _set_table_width(t_title)
    _section_row(t_title, 0, "Dichiarazione di necessità auto sostitutiva", 1,
                 row_height=380)
    t_title.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Tabella 1: dati intestazione (4 col) ---
    t1 = doc.add_table(rows=3, cols=4)
    t1.style = "Table Grid"
    _set_table_width(t1)

    # riga 0: riferimento vettura (targa cliente + marca)
    rif = " — ".join(filter(None, [data.cliente_targa, data.cliente_marca]))
    c = _merge_row(t1, 0, 0, 3)
    _lv(c, "Riferimento vettura", rif)
    _set_row_height(t1, 0, DATA_ROW_H)

    # riga 1: sottoscritto | assicurato con
    c = _merge_row(t1, 1, 0, 1); _lv(c, "Sottoscritto", data.locatario_nome)
    c = _merge_row(t1, 1, 2, 3); _lv(c, "Assicurato con", data.dich_assicurazione)
    _set_row_height(t1, 1, DATA_ROW_H)

    # riga 2: polizza
    c = _merge_row(t1, 2, 0, 3); _lv(c, "Polizza n°", data.dich_polizza)
    _set_row_height(t1, 2, DATA_ROW_H)

    # --- Tabella 2: proprietario veicolo (3 col) ---
    t2 = doc.add_table(rows=2, cols=3)
    t2.style = "Table Grid"
    _set_table_width(t2)

    _section_row(t2, 0, "Proprietario del veicolo", 3)
    _lv(t2.rows[1].cells[0], "Marca",   data.cliente_marca)
    _lv(t2.rows[1].cells[1], "Modello", data.cliente_modello)
    _lv(t2.rows[1].cells[2], "Targa",   data.cliente_targa)
    _set_row_height(t2, 1, DATA_ROW_H)

    # --- Tabella 3: dichiarazione sinistro + motivazione (2 col) ---
    n_mot = len(MOTIVAZIONI)
    t3 = doc.add_table(rows=2 + n_mot, cols=2)
    t3.style = "Table Grid"
    _set_table_width(t3)

    _section_row(t3, 0, "Dichiara che", 2)

    # riga 1: testo sinistro + data
    c = _merge_row(t3, 1, 0, 1)
    _set_cell_margins(c)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    _run(p, "In conseguenza del sinistro avvenuto il  ", size=9.0)
    _run(p, data.dich_data_sinistro or "_______________", bold=bool(data.dich_data_sinistro),
         size=9.0)
    _run(p, ",  il mezzo di sua proprietà risulta inutilizzabile e che necessita "
            "di veicolo sostitutivo per la seguente motivazione:", size=9.0)
    _set_row_height(t3, 1, DATA_ROW_H + 120)

    # righe motivazioni: 2 col — simbolo | testo
    for i, (key, label) in enumerate(MOTIVAZIONI, start=2):
        selected = (data.dich_motivazione == key)
        mark = "●" if selected else "○"
        # col 0: simbolo (stretto)
        c_mark = t3.rows[i].cells[0]
        _set_cell_margins(c_mark, left=120, right=40)
        c_mark.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _run(c_mark.paragraphs[0], mark, bold=selected, size=11.0,
             color=COLOR_WHITE if selected else None)
        if selected:
            _set_cell_shading(c_mark, COLOR_HEADER_BG)
        # col 1: testo
        c_txt = t3.rows[i].cells[1]
        _set_cell_margins(c_txt)
        c_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _run(c_txt.paragraphs[0], label, bold=selected, size=9.0)
        _set_row_height(t3, i, 320)

    # --- Testo fisso (paragrafo, italic come disclaimer danni) ---
    p = _para(doc, space_before=4, space_after=4)
    _run(p, "Dichiara altresì che il mezzo fornito dalla LYS AUTO S.r.l verrà utilizzato "
            "per il tempo strettamente necessario per l'esecuzione delle riparazioni.",
         italic=True, size=8.5)

    # --- Tabella 4: luogo, data, firma ---
    data_firma = data.data_ora.split(" ")[0] if data.data_ora else ""
    t4 = doc.add_table(rows=2, cols=3)
    t4.style = "Table Grid"
    _set_table_width(t4)
    _force_table_borders(t4)

    _col_header_row(t4, 0, ["Luogo", "Data", "In fede — Firma del Locatario"],
                    row_height=300)

    _set_row_height(t4, 1, 1000)
    for cell, val in [(t4.rows[1].cells[0], data.dich_luogo),
                      (t4.rows[1].cells[1], data_firma)]:
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, val, size=10.0)
    cell_firma = t4.rows[1].cells[2]
    _set_cell_margins(cell_firma)
    cell_firma.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p = cell_firma.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "........................", size=9.0)


def _add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(
        fp,
        f"{CARROZZERIA_NOME}  ·  {CARROZZERIA_VIA}, {CARROZZERIA_CAP} {CARROZZERIA_COMUNE}"
        f"  ·  P.IVA {CARROZZERIA_PIVA}",
        size=7.5, color=COLOR_GREY,
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def generate(data: VerbaleData, out: BinaryIO | None = None) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    _set_page_margins(doc)
    _add_header(doc, data)
    _add_locatario(doc, data)
    _add_veicolo(doc, data)

    if data.tipo == TIPO_USCITA:
        _add_franchigie(doc, data)

    _add_danni(doc, data)
    _add_note(doc, data)
    _add_firme(doc, data)

    if data.tipo == TIPO_USCITA:
        _add_dichiarazione(doc, data)

    _add_footer(doc)

    buf = BytesIO()
    doc.save(buf)
    raw = buf.getvalue()
    if out is not None:
        out.write(raw)
    return raw


def filename_for(data: VerbaleData) -> str:
    targa_safe = "".join(c for c in (data.targa or "XX") if c.isalnum())
    from datetime import date
    today = date.today().strftime("%Y%m%d")
    return f"Verbale_{data.label_tipo}_{data.numero_pratica}_{targa_safe}_{today}.docx"
