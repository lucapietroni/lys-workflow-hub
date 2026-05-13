"""Generatore del documento Word di Cessione del Credito.

Costruisce il file `.docx` da zero usando python-docx, con un layout pulito che
sostituisce quello originale (un po' grezzo) ma mantiene il testo legale
identico parola per parola.

Uso:
    from io import BytesIO
    data = from_pratica(pratica)  # vedi data.py
    buf = BytesIO()
    generate(data, buf)
    pdf_bytes = buf.getvalue()
"""
from __future__ import annotations

from datetime import date
from io import BytesIO
from typing import BinaryIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from lys_workflow_hub.workflows.cessione_credito.data import (
    CARROZZERIA_COMUNE,
    CARROZZERIA_NOME,
    CARROZZERIA_PIVA,
    CARROZZERIA_VIA,
    CARROZZERIA_CAP,
    CessioneData,
    LUOGO_SOTTOSCRIZIONE,
)


# Palette
COLOR_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
COLOR_GREY = RGBColor(0x4A, 0x55, 0x68)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)


# -----------------------------------------------------------------------------
# Helpers di basso livello su python-docx
# -----------------------------------------------------------------------------


def _set_default_font(doc, name: str = "Calibri", size_pt: float = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
             color: RGBColor | None = None, size: float | None = None) -> None:
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def _add_segments(paragraph, segments: list[tuple[str, bool]]) -> None:
    """Aggiunge piu' run nello stesso paragrafo. Ogni segmento e' (testo, bold)."""
    for text, bold in segments:
        if text:
            _add_run(paragraph, text, bold=bold)


def _add_title(doc, text: str) -> None:
    """Titolo grande centrato con linea sotto."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _add_run(p, text, bold=True, size=22, color=COLOR_BLACK)
    # Linea decorativa sottile
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(18)
    _add_run(p2, "─" * 32, color=COLOR_BLACK)


def _add_section_label(doc, text: str) -> None:
    """Etichetta di sezione (PREMESSO, DICHIARA)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    _add_run(p, text, bold=True, size=13, color=COLOR_BLACK)


def _justified(doc):
    """Crea un paragrafo giustificato con spaziatura standard."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def _set_page_margins(doc) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_signature_table(doc, *, with_date: str | None = None) -> None:
    """Tabella firme a due colonne: Cessionario | Cedente.

    Se `with_date` e' valorizzato, una riga sopra ospita "Roma, <data>".
    """
    if with_date:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        _add_run(p, with_date, italic=True)

    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    for col in table.columns:
        col.width = Cm(8.0)
    # Riga 1: etichette
    labels = table.rows[0].cells
    for cell, label in zip(labels, ["Firma Cessionario", "Firma Cedente"]):
        cell.width = Cm(8.0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        _add_run(para, label, bold=True, color=COLOR_BLACK, size=10)
        para.paragraph_format.space_before = Pt(36)  # spazio per la firma sopra
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Linea sopra (top border della cella)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "6")
        top.set(qn("w:space"), "0")
        top.set(qn("w:color"), "000000")
        tc_borders.append(top)
        for side in ("left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "nil")
            tc_borders.append(b)
        tc_pr.append(tc_borders)
    # Riga 2: vuota per spaziatura
    for cell in table.rows[1].cells:
        cell.text = ""


# -----------------------------------------------------------------------------
# Generatore principale
# -----------------------------------------------------------------------------


def generate(data: CessioneData, out: BinaryIO | None = None) -> bytes:
    """Costruisce il documento Word di cessione del credito.

    Se `out` e' fornito, ci scrive dentro. In ogni caso restituisce i byte
    del documento generato (utile per i test e per il download HTTP).
    """
    doc = Document()
    _set_default_font(doc)
    _set_page_margins(doc)

    # ------------- Titolo -------------
    _add_title(doc, "CESSIONE DI CREDITO")

    # ------------- Blocco cedente (anagrafica) -------------
    p = _justified(doc)
    _add_segments(p, [
        (f"{data.sottoscritto_label} ", False),
        (data.cedente_nome_completo, True),
        (f", {data.nato_label} ", False),
        (data.cedente_luogo_nascita, True),
        (" il ", False),
        (data.cedente_data_nascita_formattata, True),
        (", residente in ", False),
        (data.cedente_residenza_via, True),
        (", ", False),
        (data.cedente_residenza_comune, True),
        (", codice fiscale ", False),
        (data.cedente_codice_fiscale, True),
        (
            (
                f", in qualità di legale rappresentante pro-tempore della ditta "
            )
            if data.e_ditta else "",
            False,
        ),
        (data.ditta_nome if data.e_ditta else "", True),
        (", partita IVA " if data.e_ditta else "", False),
        (data.ditta_partita_iva if data.e_ditta else "", True),
        (" (cedente).", False),
    ])

    # ------------- PREMESSO -------------
    _add_section_label(doc, "PREMESSO")

    p = _justified(doc)
    _add_segments(p, [
        ("che in data ", False),
        (data.sinistro_data_formattata, True),
        (" alle ore ", False),
        (data.sinistro_ora, True),
        (" nel Comune di ", False),
        (data.sinistro_comune, True),
        (", Via ", False),
        (data.sinistro_via, True),
        (
            ", ha subito un incidente stradale per la collisione tra il veicolo "
            "di sua proprietà ",
            False,
        ),
        (data.veicolo_cedente_descrizione, True),
        (" targato ", False),
        (data.veicolo_cedente_targa, True),
        (" ed il veicolo ", False),
        (data.controparte_veicolo_descrizione, True),
        (", targato ", False),
        (data.controparte_veicolo_targa, True),
        (", di proprietà del Sig./ra/ Ditta ", False),
        (data.controparte_proprietario, True),
        (" condotto dal Sig/Sig.ra ", False),
        (data.controparte_conducente, True),
        (";", False),
    ])

    p = _justified(doc)
    _add_segments(p, [
        (
            "che l'esponente ritiene che la responsabilità vada ascritta alla "
            "controparte sopra indicata, assicurata per la responsabilità civile "
            "con la compagnia ",
            False,
        ),
        (data.controparte_compagnia, True),
        (", polizza n. ", False),
        (data.controparte_polizza, True),
        (", in quanto: ", False),
        (data.sinistro_dinamica, True),
        (";", False),
    ])

    p = _justified(doc)
    _add_run(
        p,
        "che la dinamica dell'incidente si evince dalla CAI che si allega e/o dal "
        "verbale di incidente redatto dalle autorità intervenute, che il danneggiato "
        "si impegna a consegnare alla carrozzeria cessionaria;",
    )

    p = _justified(doc)
    _add_segments(p, [
        (
            "che il sottoscritto intende effettuare le riparazioni del proprio "
            "veicolo presso la ",
            False,
        ),
        (CARROZZERIA_NOME, True),
        (", sita in ", False),
        (CARROZZERIA_VIA, True),
        (", Comune di ", False),
        (CARROZZERIA_COMUNE, True),
        (".", False),
    ])

    p = _justified(doc)
    _add_run(p, "Tutto ciò premesso,", italic=True)

    # ------------- DICHIARA -------------
    _add_section_label(doc, "DICHIARA")

    # Clausola 1 - cessione vera e propria. Testo originale invariato.
    p = _justified(doc)
    _add_segments(p, [
        (
            "di cedere ai sensi e per gli effetti degli artt. 1260 sgg. cod. civ. il "
            "credito per il risarcimento dei danni materiali derivanti dal sinistro "
            "di cui sopra e per il rimborso delle somme corrisposte per il noleggio "
            "di veicolo sostitutivo, nonché il danno da fermo tecnico, alla ",
            False,
        ),
        (CARROZZERIA_NOME, True),
        (
            " in persona del legale rappresentante pro-tempore (cessionaria), con sede "
            "legale nel Comune di ",
            False,
        ),
        (CARROZZERIA_COMUNE, True),
        (", ", False),
        (CARROZZERIA_VIA, True),
        (", ", False),
        (CARROZZERIA_CAP, True),
        (", P.IVA ", False),
        (CARROZZERIA_PIVA, True),
        (
            ", nonché ogni altra voce di danno patrimoniale connessa e consequenziale "
            "al sinistro, fino alla concorrenza del valore complessivo delle fatture "
            "che saranno emesse dalla Cessionaria per le prestazioni di riparazione e "
            "i servizi accessori forniti;",
            False,
        ),
    ])

    # Clausole successive - testo originale invariato.
    altre_clausole = [
        "La presente cessione del credito de quo è espressamente irrevocabile e "
        "l'obbligazione originaria del Cedente si estinguerà unicamente con "
        "l'effettiva riscossione da parte della Cessionaria del credito ceduto;",

        "ai sensi dell'art. 1266 c.c., il Cedente garantisce l'esistenza, la "
        "validità e la titolarità del credito al momento della cessione, nonché "
        "dichiara che il credito non è stato precedentemente ceduto a terzi;",

        "ove, per qualsivoglia motivo, si manifestasse l'impossibilità di recupero "
        "del predetto importo nei confronti del responsabile civile e/o della sua "
        "compagnia di assicurazione e/o della compagnia di assicurazione del "
        "danneggiato ai sensi dell'art. 149 C.d.A., il sottoscritto si impegna a "
        f"rimborsare il predetto importo alla {CARROZZERIA_NOME} ai sensi dell'art. "
        "1267 c.c., a semplice richiesta di quest'ultima, senza alcuna corresponsione "
        "di interesse o altro;",

        "ove peraltro, ricevuta la richiesta di rimborso del cessionario, il "
        "pagamento non fosse effettuato dal sottoscritto cedente entro 30 (trenta) "
        "giorni dalla richiesta, saranno dovuti anche gli interessi di mora al tasso "
        "legale nonché la rivalutazione monetaria;",

        "il sottoscritto cedente, nel contempo, si obbliga a corrispondere alla "
        f"{CARROZZERIA_NOME} il costo della riparazione non corrisposto dai "
        "responsabili, in quanto superiore al valore commerciale ante-sinistro del "
        "mezzo;",

        "inoltre il sottoscritto cedente si obbliga a conferire il mandato di "
        f"gestione della pratica per il recupero dei danni (conseguenti al prefatto "
        f"sinistro) alla {CARROZZERIA_NOME} in forma irrevocabile.",
    ]
    for testo in altre_clausole:
        p = _justified(doc)
        _add_run(p, testo)

    # ------------- Firme (prima coppia) -------------
    _add_signature_table(doc)

    # ------------- Clausole vessatorie (art. 1341-1342) -------------
    p = _justified(doc)
    p.paragraph_format.space_before = Pt(18)
    _add_run(
        p,
        "Ai sensi e per gli effetti degli artt. 1341 e 1342 c.c. il cedente sottoscrive "
        "per accettazione e ratifica le seguenti clausole: 2) irrevocabilità della "
        "cessione; 3) garanzia sull'esistenza del credito; 4) garanzia di solvenza "
        "dell'obbligazione ex art. 1267 c.c.; 5) termini di pagamento e interessi di "
        "mora; 7) mandato per la gestione del sinistro.",
    )

    # Roma, data firma + seconda coppia di firme
    oggi = date.today().strftime("%d/%m/%Y")
    _add_signature_table(doc, with_date=f"{LUOGO_SOTTOSCRIZIONE}, {oggi}")

    # ------------- Footer -------------
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        fp,
        f"{CARROZZERIA_NOME}  ·  {CARROZZERIA_VIA}, {CARROZZERIA_CAP} {CARROZZERIA_COMUNE}  ·  P.IVA {CARROZZERIA_PIVA}",
        color=COLOR_BLACK, size=8,
    )

    # ------------- Output -------------
    buf = BytesIO()
    doc.save(buf)
    raw = buf.getvalue()
    if out is not None:
        out.write(raw)
    return raw


def filename_for(data: CessioneData) -> str:
    """Nome file standard per il documento generato."""
    today = date.today().strftime("%Y%m%d")
    # Compatti il nome: spazi -> _, solo caratteri amici
    nome_safe = "".join(c if c.isalnum() else "_" for c in data.cedente_nome_completo)
    nome_safe = nome_safe.strip("_") or "cedente"
    return f"Cessione_credito_{data.numero_pratica}_{nome_safe}_{today}.docx"
